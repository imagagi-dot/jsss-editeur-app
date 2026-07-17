#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
remplir_article.py — Injecte le contenu d'un manuscrit (JSON) dans la COQUILLE
reelle clonee de la maquette JSSS (TEMPLATE_article_JSSS.docx) pour produire un
article typographie conforme (bandeau logos, titres bleus, tableau dates+resume,
corps 2 colonnes, encadre affiliations, references).

Balisage en ligne : **gras**  et  ^exposant^
Usage : python remplir_article.py article.json sortie.docx [chemin_template]
"""
import copy
import json
import re
import sys
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

F_BODY = "Palatino Linotype"
SZ = dict(section=12, body=10, sub=10, caption=9, ref=9)
TOKEN = re.compile(r"(\*\*.+?\*\*|\*[^\*]+\*|\^[^\^]+\^)")
DEFAULT_TEMPLATE = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                                "TEMPLATE_article_JSSS.docx")


def _clone_rpr(run_el):
    rpr = run_el.find(qn("w:rPr")) if run_el is not None else None
    return copy.deepcopy(rpr) if rpr is not None else None


def _mk_run(rpr_template, content, bold=None, italic=None, sup=False):
    r = OxmlElement("w:r")
    if rpr_template is not None:
        r.append(copy.deepcopy(rpr_template))
    rpr = r.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        r.insert(0, rpr)
    if bold is not None:
        for tag in ("w:b", "w:bCs"):
            for e in rpr.findall(qn(tag)):
                rpr.remove(e)
        if bold:
            rpr.append(OxmlElement("w:b"))
    if italic is not None:
        for tag in ("w:i", "w:iCs"):
            for e in rpr.findall(qn(tag)):
                rpr.remove(e)
        if italic:
            rpr.append(OxmlElement("w:i"))
    if sup:
        va = OxmlElement("w:vertAlign")
        va.set(qn("w:val"), "superscript")
        rpr.append(va)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = content
    r.append(t)
    return r


def fill_paragraph(paragraph, text, base_bold=None, base_italic=None):
    p = paragraph._p
    runs = p.findall(qn("w:r"))
    rpr_tmpl = _clone_rpr(runs[0]) if runs else None
    for r in runs:
        p.remove(r)
    for part in TOKEN.split(text):
        if not part:
            continue
        b, it, sup, content = base_bold, base_italic, False, part
        if part.startswith("**") and part.endswith("**"):
            b, content = True, part[2:-2]
        elif part.startswith("*") and part.endswith("*"):
            it, content = True, part[1:-1]
        elif part.startswith("^") and part.endswith("^"):
            sup, content = True, part[1:-1]
        p.append(_mk_run(rpr_tmpl, content, bold=b, italic=it, sup=sup))
    return paragraph


def fill_cell(cell, items):
    tc = cell._tc
    paras = tc.findall(qn("w:p"))
    model = paras[0]
    r0 = model.find(qn("w:r"))
    rpr_tmpl = _clone_rpr(r0) if r0 is not None else None
    for p in paras[1:]:
        tc.remove(p)
    base = model
    for r in base.findall(qn("w:r")):
        base.remove(r)
    if isinstance(items, str):
        items = [items]

    def render_into(p_el, text, subhead=None):
        if subhead:
            p_el.append(_mk_run(rpr_tmpl, subhead + " : ", bold=True))
            for part in TOKEN.split(text):
                if not part:
                    continue
                b, it, sup, content = False, False, False, part
                if part.startswith("**") and part.endswith("**"):
                    b, content = True, part[2:-2]
                elif part.startswith("*") and part.endswith("*"):
                    it, content = True, part[1:-1]
                elif part.startswith("^") and part.endswith("^"):
                    sup, content = True, part[1:-1]
                p_el.append(_mk_run(rpr_tmpl, content, bold=b, italic=it, sup=sup))
        else:
            for part in TOKEN.split(text):
                if not part:
                    continue
                b, it, sup, content = False, False, False, part
                if part.startswith("**") and part.endswith("**"):
                    b, content = True, part[2:-2]
                elif part.startswith("*") and part.endswith("*"):
                    it, content = True, part[1:-1]
                elif part.startswith("^") and part.endswith("^"):
                    sup, content = True, part[1:-1]
                p_el.append(_mk_run(rpr_tmpl, content, bold=b, italic=it, sup=sup))

    first = True
    for it in items:
        if first:
            p_el = base
            first = False
        else:
            p_el = copy.deepcopy(model)
            for r in p_el.findall(qn("w:r")):
                p_el.remove(r)
            base.addnext(p_el)
            base = p_el
        if isinstance(it, dict):
            render_into(p_el, it.get("t", ""), subhead=it.get("h"))
        else:
            render_into(p_el, it)


def _new_para(text, size, bold=False, italic=False, color=None, hanging_indent=False):
    p = OxmlElement("w:p")
    ppr = OxmlElement("w:pPr")
    jc = OxmlElement("w:jc"); jc.set(qn("w:val"), "both"); ppr.append(jc)
    sp = OxmlElement("w:spacing"); sp.set(qn("w:after"), "80"); ppr.append(sp)
    if hanging_indent:
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "360")
        ind.set(qn("w:hanging"), "360")
        ppr.append(ind)
    p.append(ppr)
    para = Paragraph(p, None)
    for part in TOKEN.split(text):
        if not part:
            continue
        b, sup, it, content = bold, False, italic, part
        if part.startswith("**") and part.endswith("**"):
            b, content = True, part[2:-2]
        elif part.startswith("^") and part.endswith("^"):
            sup, content = True, part[1:-1]
        run = para.add_run(content)
        run.font.name = F_BODY
        run.font.size = Pt(size)
        run.font.bold = b
        run.font.italic = it
        if sup:
            run.font.superscript = True
        if color:
            from docx.shared import RGBColor
            run.font.color.rgb = RGBColor(*color)
    return p


def _figure_element(doc, image_path, width_cm=7.5):
    from docx.enum.text import WD_ALIGN_PARAGRAPH as A
    p = doc.add_paragraph()
    p.alignment = A.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(image_path, width=Cm(width_cm))
    el = p._p
    el.getparent().remove(el)
    return el

def _table_element(doc, table_data):
    if not table_data or not table_data[0]:
        return None
    rows = len(table_data)
    cols = len(table_data[0])
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    for i, row in enumerate(table_data):
        for j, cell_text in enumerate(row):
            if j < len(table.rows[i].cells):
                table.cell(i, j).text = str(cell_text)
    
    el = table._element
    el.getparent().remove(el)
    return el


def build_body_elements(spec, doc=None, images_dir=""):
    els = []
    for blk in spec.get("body", []):
        t = blk.get("type", "p"); txt = blk.get("text", "")
        if t == "h2":
            els.append(_new_para(txt, SZ["section"], bold=True))
        elif t == "h3":
            els.append(_new_para(txt, SZ["sub"], bold=True))
        elif t == "caption":
            els.append(_new_para(txt, SZ["caption"], bold=True))
        elif t == "figure":
            img = blk.get("image", "")
            if images_dir and img and not os.path.isabs(img):
                img = os.path.join(images_dir, img)
            if doc is not None and img and os.path.exists(img):
                try:
                    els.append(_figure_element(doc, img, blk.get("width_cm", 7.5)))
                except Exception as e:
                    els.append(_new_para("[A COMPLETER : image %s (%s)]" % (os.path.basename(img), e), SZ["body"]))
            else:
                els.append(_new_para("[A COMPLETER : inserer l'image %s]" % (os.path.basename(img) or "?"), SZ["body"]))
            if blk.get("caption"):
                els.append(_new_para(blk["caption"], SZ["caption"], bold=True))
        elif t == "table":
            if doc is not None and "data" in blk:
                tbl_el = _table_element(doc, blk["data"])
                if tbl_el is not None:
                    els.append(tbl_el)
            if blk.get("caption"):
                els.append(_new_para(blk["caption"], SZ["caption"], bold=True))
        else:
            els.append(_new_para(txt, SZ["body"]))
    conflict = spec.get("conflict", "Les auteurs ne declarent aucun conflit d'interets.")
    els.append(_new_para("**Conflit d'interets :** " + conflict, SZ["body"]))
    if spec.get("references"):
        els.append(_new_para("References", SZ["section"], bold=True))
        for i, ref in enumerate(spec["references"], 1):
            els.append(_new_para("%d. %s" % (i, ref), SZ["ref"], color=(0, 0, 255), hanging_indent=True))
    return els


def set_running_header(doc, citation):
    if not citation:
        return
    line = citation + "\t\tISSN : 1859-5162"
    for sec in doc.sections:
        for hdr in (sec.first_page_header, sec.header, sec.even_page_header):
            if hdr is None:
                continue
            try:
                hdr.is_linked_to_previous = False
            except Exception:
                pass
            if not hdr.paragraphs:
                hdr.add_paragraph()
            p = hdr.paragraphs[0]
            p.clear()
            
            # Éliminer le formatage direct "sale" du paragraphe (qui forçait Calibri)
            from docx.oxml.ns import qn
            pPr = p._element.pPr
            if pPr is not None:
                rPr_in_p = pPr.find(qn('w:rPr'))
                if rPr_in_p is not None:
                    pPr.remove(rPr_in_p)
            
            # Modification du style global de l'entête
            from docx.shared import Pt
            try:
                style = doc.styles['Header']
            except KeyError:
                try:
                    style = doc.styles['En-tête']
                except KeyError:
                    style = None
            
            if style:
                font = style.font
                font.name = "Trebuchet MS"
                font.size = Pt(8)
                font.bold = True
                font.italic = True
            
            p.style = style
            
            parts = line.split('\t')
            for i, part in enumerate(parts):
                if part:
                    run = p.add_run(part)
                    run.font.name = "Trebuchet MS"
                    run.font.size = Pt(8)
                    run.font.bold = True
                    run.font.italic = True
                if i < len(parts) - 1:
                    run_tab = p.add_run()
                    run_tab.add_tab()


def fill_affiliation_box(box_p, aff_lines, corr):
    txbx = box_p.find(".//" + qn("w:txbxContent"))
    if txbx is None:
        return
    tb_paras = txbx.findall(".//" + qn("w:p"))
    if not tb_paras:
        return
    model = tb_paras[1] if len(tb_paras) > 1 else tb_paras[0]
    corr_idx = None
    for i, p in enumerate(tb_paras):
        tt = "".join(n.text or "" for n in p.iter(qn("w:t")))
        if "correspondant" in tt.lower():
            corr_idx = i
            break
    end = corr_idx if corr_idx is not None else len(tb_paras)
    old_lines = tb_paras[1:end]
    parent = model.getparent()
    for p in old_lines:
        parent.remove(p)
    prev = tb_paras[0]
    rpr = _clone_rpr(model.find(qn("w:r")))
    for line in aff_lines:
        line = re.sub(r"^\s*\d+\s*[-.)]\s*", "", line)
        newp = copy.deepcopy(model)
        for r in newp.findall(qn("w:r")):
            newp.remove(r)
        
        for part in TOKEN.split(line):
            if not part:
                continue
            b, it, sup, content = False, False, False, part
            if part.startswith("**") and part.endswith("**"):
                b, content = True, part[2:-2]
            elif part.startswith("*") and part.endswith("*"):
                it, content = True, part[1:-1]
            elif part.startswith("^") and part.endswith("^"):
                sup, content = True, part[1:-1]
            newp.append(_mk_run(rpr, content, bold=b, italic=it, sup=sup))

        prev.addnext(newp)
        prev = newp
    if corr_idx is not None and corr:
        cp = tb_paras[corr_idx]
        rc = _clone_rpr(cp.find(qn("w:r")))
        for tag in ("w:r", "w:hyperlink"):
            for e in cp.findall(qn(tag)):
                cp.remove(e)
        cp.append(_mk_run(rc, "*Auteur correspondant : ", bold=True))
        cp.append(_mk_run(rc, corr))


def fill(spec, template, out_path):
    doc = Document(template)
    body = doc.element.body

    for p in doc.paragraphs:
        s = p.text.strip()
        if s.startswith("Indications et complications des cholecystectomies") or \
           s.startswith("Indications et complications des chol"):
            fill_paragraph(p, spec.get("title_fr", "[A COMPLETER : titre]"))
        elif s.startswith("Indications and complications of open cholecystectomy") or \
             s.startswith("Indications and complications of open chol"):
            fill_paragraph(p, spec.get("title_en", "[A COMPLETER : titre EN]"), base_italic=True)
        elif s.startswith("Mamoudou Hamidou AN"):
            fill_paragraph(p, spec.get("authors", "[A COMPLETER : auteurs]"))

    t = doc.tables[0]
    if spec.get("date_recu"):
        fill_paragraph(t.rows[0].cells[0].paragraphs[0], spec["date_recu"])
    if spec.get("date_accepte"):
        fill_paragraph(t.rows[0].cells[3].paragraphs[0], spec["date_accepte"])
    if spec.get("date_publie"):
        fill_paragraph(t.rows[0].cells[4].paragraphs[0], spec["date_publie"])
    kw_str = str(spec.get('keywords', '[A COMPLETER]')).strip()
    kw_str = re.sub(r'^(?:[\*\s]*(?:Mots-cl[eé]s?|Keywords)[\*\s:]*)+', '', kw_str, flags=re.IGNORECASE).strip()
    kw_fr = [f"*{kw_str}*"]
    
    resume_list = spec.get("resume") if isinstance(spec.get("resume"), list) else [{"t": str(spec.get("resume", "[A COMPLETER : resume]"))}]
    if resume_list:
        first_h = resume_list[0].get("h", "").strip()
        first_t = resume_list[0].get("t", "").strip()
        first_h_clean = re.sub(r'^(?:[\*\s]*(?:R[eé]sum[eé])[\*\s:]*)+$', '', first_h, flags=re.IGNORECASE).strip()
        first_t_clean = re.sub(r'^(?:[\*\s]*(?:R[eé]sum[eé])[\*\s:]*)+', '', first_t, flags=re.IGNORECASE).strip()
        if first_h_clean == "" and first_t_clean == "":
            resume_list = resume_list[1:]
        else:
            resume_list[0] = {"h": first_h_clean, "t": first_t_clean}
    res_fr = resume_list

    kw_en_str = str(spec.get('keywords_en', '[TO COMPLETE]')).strip()
    kw_en_str = re.sub(r'^(?:[\*\s]*(?:Keywords|Mots-cl[eé]s?)[\*\s:]*)+', '', kw_en_str, flags=re.IGNORECASE).strip()
    kw_en = [f"*{kw_en_str}*"]
    
    abs_list = spec.get("abstract") if isinstance(spec.get("abstract"), list) else [{"t": str(spec.get("abstract", "[TO COMPLETE: abstract]"))}]
    if abs_list:
        first_h = abs_list[0].get("h", "").strip()
        first_t = abs_list[0].get("t", "").strip()
        first_h_clean = re.sub(r'^(?:[\*\s]*(?:Abstract)[\*\s:]*)+$', '', first_h, flags=re.IGNORECASE).strip()
        first_t_clean = re.sub(r'^(?:[\*\s]*(?:Abstract)[\*\s:]*)+', '', first_t, flags=re.IGNORECASE).strip()
        if first_h_clean == "" and first_t_clean == "":
            abs_list = abs_list[1:]
        else:
            abs_list[0] = {"h": first_h_clean, "t": first_t_clean}
    res_en = abs_list

    fill_cell(t.rows[2].cells[0], kw_fr)
    fill_cell(t.rows[2].cells[1], res_fr)
    fill_cell(t.rows[4].cells[0], kw_en)
    fill_cell(t.rows[4].cells[2], res_en)

    aff_lines = spec.get("affiliations", [])
    corr = spec.get("corresponding", "")
    aff_box_p = None
    for el in body:
        if el.tag != qn("w:p"):
            continue
        txbx = el.find(".//" + qn("w:txbxContent"))
        if txbx is not None:
            inner = "".join(n.text or "" for n in txbx.iter(qn("w:t")))
            if "Affiliation" in inner:
                if not aff_lines and not corr:
                    body.remove(el)
                else:
                    aff_box_p = el
                    fill_affiliation_box(el, aff_lines, corr)
                break

    kids = list(body)
    tbl_el = t._tbl
    idx_tbl = kids.index(tbl_el)
    final_sectPr = kids[-1] if kids[-1].tag == qn("w:sectPr") else None
    sect_break_p = None
    to_remove = []
    for el in kids[idx_tbl + 1:]:
        if el is final_sectPr:
            continue
        if el is aff_box_p:
            continue
        ppr = el.find(qn("w:pPr")) if el.tag == qn("w:p") else None
        if ppr is not None and ppr.find(qn("w:sectPr")) is not None and sect_break_p is None:
            sect_break_p = el
            continue
        to_remove.append(el)
    for el in to_remove:
        body.remove(el)

    if final_sectPr is not None:
        cols = final_sectPr.find(qn("w:cols"))
        if cols is None:
            cols = OxmlElement("w:cols"); final_sectPr.append(cols)
        cols.set(qn("w:num"), "2")
        cols.set(qn("w:space"), "708")

    images_dir = spec.get("images_dir", "")
    anchor = final_sectPr
    for el in build_body_elements(spec, doc=doc, images_dir=images_dir):
        if anchor is not None:
            anchor.addprevious(el)
        else:
            body.append(el)

    set_running_header(doc, spec.get("header_citation"))
    doc.save(out_path)
    print("OK -> " + out_path)


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python remplir_article.py article.json sortie.docx [template]")
        sys.exit(1)
    tpl = sys.argv[3] if len(sys.argv) > 3 else DEFAULT_TEMPLATE
    with open(sys.argv[1], "r", encoding="utf-8") as fh:
        spec = json.load(fh)
    fill(spec, tpl, sys.argv[2])
